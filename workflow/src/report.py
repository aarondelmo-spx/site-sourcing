"""Word document report generator for the SPX weekly Ops & Compliance report."""
import os
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── Colour palette ──────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1F, 0x38, 0x64)
TEAL   = RGBColor(0x1B, 0x6B, 0x93)
GREY   = RGBColor(0x88, 0x88, 0x88)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x00, 0x00, 0x00)

RAG_COLORS = {
    "RED":   RGBColor(0xFF, 0xCC, 0xCC),
    "AMBER": RGBColor(0xFF, 0xF0, 0xB0),
    "GREEN": RGBColor(0xCC, 0xFF, 0xCC),
}
RAG_TEXT = {
    "RED":   RGBColor(0xCC, 0x00, 0x00),
    "AMBER": RGBColor(0x99, 0x66, 0x00),
    "GREEN": RGBColor(0x00, 0x77, 0x00),
}


# ─── XML helpers ─────────────────────────────────────────────────────────────
def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, color="CCCCCC"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _set_row_bg(row, hex_color: str):
    for cell in row.cells:
        _set_cell_bg(cell, hex_color)


def _rgb_to_hex(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


# ─── Style helpers ───────────────────────────────────────────────────────────
def _para(cell, text, bold=False, color=BLACK, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Garamond"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def _add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Garamond"
    run.font.size = Pt(16 if level == 1 else 13)
    run.font.color.rgb = NAVY if level == 1 else TEAL
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "1F3864")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def _add_body(doc, text, bold=False, color=BLACK, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Garamond"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def _table_header_row(table, headers, widths_cm, bg="1B6B93"):
    row = table.rows[0]
    for i, (h, w) in enumerate(zip(headers, widths_cm)):
        cell = row.cells[i]
        cell.width = Cm(w)
        _set_cell_bg(cell, bg)
        _set_cell_borders(cell, "1B6B93")
        _para(cell, h, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)


def _add_table_row(table, values, bgs=None, bolds=None, aligns=None, size=9):
    row = table.add_row()
    bgs    = bgs    or ["FFFFFF"] * len(values)
    bolds  = bolds  or [False]   * len(values)
    aligns = aligns or [WD_ALIGN_PARAGRAPH.LEFT] * len(values)
    for i, (val, bg, bold, align) in enumerate(zip(values, bgs, bolds, aligns)):
        cell = row.cells[i]
        _set_cell_bg(cell, bg)
        _set_cell_borders(cell)
        _para(cell, str(val), bold=bold, size=size, align=align)
    return row


# ─── Main generator ──────────────────────────────────────────────────────────
def generate_report(
    kpi_report: dict,
    compliance_report: dict,
    hse_report: dict,
    output_path: str,
    as_of: date = None,
):
    """Generate the SPX Weekly Ops & Compliance Word report.

    Raises:
        OSError: if the output directory does not exist
        Exception: for any doc generation failure
    """
    output_dir = os.path.dirname(os.path.abspath(output_path))
    if not os.path.exists(output_dir):
        raise OSError(f"Output directory does not exist: {output_dir}")

    if as_of is None:
        as_of = date.today()

    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Cover block ──────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_r = title_p.add_run("SPX Weekly Ops & Compliance Report")
    title_r.bold = True
    title_r.font.name = "Garamond"
    title_r.font.size = Pt(22)
    title_r.font.color.rgb = NAVY

    sub_p = doc.add_paragraph()
    sub_r = sub_p.add_run(f"Week of {as_of.strftime('%d %B %Y')}  ·  Auto-generated by Cowork")
    sub_r.font.name = "Garamond"
    sub_r.font.size = Pt(11)
    sub_r.font.color.rgb = GREY

    doc.add_paragraph()  # spacer

    # ── Summary metrics bar ──────────────────────────────────────────────────
    flagged   = kpi_report["flagged_count"]
    red_count = compliance_report["red_count"]
    overdue   = hse_report["overdue_count"]
    total_inc = hse_report["total_this_week"]

    summary_table = doc.add_table(rows=1, cols=4)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    metrics = [
        (str(flagged),   "KPI Flags",        "FFEECC" if flagged > 0 else "CCFFCC"),
        (str(red_count), "Permit RED",        "FFCCCC" if red_count > 0 else "CCFFCC"),
        (str(overdue),   "Overdue Actions",   "FFCCCC" if overdue > 0 else "CCFFCC"),
        (str(total_inc), "Incidents (Week)",  "FFFFFF"),
    ]
    for i, (val, label, bg) in enumerate(metrics):
        cell = summary_table.rows[0].cells[i]
        _set_cell_bg(cell, bg)
        _set_cell_borders(cell, "CCCCCC")
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(val)
        r1.bold = True
        r1.font.name = "Garamond"
        r1.font.size = Pt(24)
        r1.font.color.rgb = NAVY
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(label)
        r2.font.name = "Garamond"
        r2.font.size = Pt(9)
        r2.font.color.rgb = GREY

    doc.add_paragraph()

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 1: KPI
    # ─────────────────────────────────────────────────────────────────────────
    _add_heading(doc, "1. KPI Performance Dashboard")
    _add_body(doc, f"Week ending {as_of.strftime('%d %b %Y')}. Threshold: flag on ±5% WoW change. "
                   f"{flagged} metric(s) flagged this week.")

    if kpi_report["top_misses"]:
        _add_heading(doc, "Top Misses", level=2)
        miss_tbl = doc.add_table(rows=1, cols=5)
        _table_header_row(miss_tbl, ["Hub", "Metric", "This Week", "Last Week", "Δ %"],
                          [3.5, 5.0, 2.5, 2.5, 2.5])
        for miss in kpi_report["top_misses"]:
            delta_str = f"{miss['delta_pct']:+.1f}%"
            rag_bg = _rgb_to_hex(RAG_COLORS[miss["rag"]])[0:6] if miss["rag"] in RAG_COLORS else "FFEECC"
            _add_table_row(miss_tbl,
                           [miss["hub"], miss["metric"],
                            f"{miss['this_week']:.1f}", f"{miss['last_week']:.1f}", delta_str],
                           bgs=["F5F5F5", "FFFFFF", "FFFFFF", "FFFFFF", "FFEECC"],
                           bolds=[True, False, False, False, True])

    _add_heading(doc, "Full KPI Matrix", level=2)
    kpi_tbl = doc.add_table(rows=1, cols=7)
    _table_header_row(kpi_tbl, ["Hub", "Metric", "This Wk", "Last Wk", "Target", "Δ %", "RAG"],
                      [3.0, 4.5, 2.0, 2.0, 2.0, 2.0, 1.5])
    for r in kpi_report["rows"]:
        rag_bg = {"RED": "FFCCCC", "AMBER": "FFF0B0", "GREEN": "CCFFCC"}.get(r["rag"], "FFFFFF")
        delta_str = f"{r['delta_pct']:+.1f}%"
        flag_bg = "FFEECC" if r["flagged"] else "FFFFFF"
        _add_table_row(kpi_tbl,
                       [r["hub"], r["metric"], f"{r['this_week']:.1f}",
                        f"{r['last_week']:.1f}", f"{r['target']:.1f}", delta_str, r["rag"]],
                       bgs=["F5F5F5", "FFFFFF", "FFFFFF", "FFFFFF", "FFFFFF", flag_bg, rag_bg],
                       bolds=[True, False, False, False, False, r["flagged"], True],
                       aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT,
                                WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                                WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                                WD_ALIGN_PARAGRAPH.CENTER])

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 2: COMPLIANCE
    # ─────────────────────────────────────────────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, "2. Compliance Permit Risk Matrix")
    _add_body(doc, (
        f"Total permits tracked: {compliance_report['total_permits']}  ·  "
        f"RED (≤30 days): {compliance_report['red_count']}  ·  "
        f"AMBER (31–60 days): {compliance_report['amber_count']}  ·  "
        f"GREEN (>60 days): {compliance_report['green_count']}"
    ))

    if compliance_report["action_list"]:
        _add_heading(doc, "Action Required (RED + AMBER)", level=2)
        act_tbl = doc.add_table(rows=1, cols=6)
        _table_header_row(act_tbl,
                          ["Hub", "Site", "Permit Type", "Expiry Date", "Days", "Status"],
                          [3.0, 2.0, 4.0, 2.5, 1.5, 3.0])
        for item in compliance_report["action_list"]:
            rag_bg = {"RED": "FFCCCC", "AMBER": "FFF0B0"}.get(item["rag"], "FFFFFF")
            days_str = str(item["days_to_expiry"]) if item["days_to_expiry"] >= 0 else f"EXPIRED {abs(item['days_to_expiry'])}d ago"
            _add_table_row(act_tbl,
                           [item["hub"], item["site_code"], item["permit_type"],
                            item["expiry_date"], days_str, item["renewal_status"]],
                           bgs=["F5F5F5", "F5F5F5", "FFFFFF", "FFFFFF", rag_bg, "FFFFFF"],
                           bolds=[True, False, False, False, True, False])

    _add_heading(doc, "Full Permit Register", level=2)
    full_tbl = doc.add_table(rows=1, cols=6)
    _table_header_row(full_tbl,
                      ["Hub", "Site", "Permit Type", "Expiry Date", "Days", "RAG"],
                      [3.0, 2.0, 4.0, 2.5, 1.5, 3.0])
    for r in sorted(compliance_report["rows"], key=lambda x: x["days_to_expiry"]):
        rag_bg = {"RED": "FFCCCC", "AMBER": "FFF0B0", "GREEN": "CCFFCC"}.get(r["rag"], "FFFFFF")
        days_str = str(r["days_to_expiry"]) if r["days_to_expiry"] >= 0 else f"EXPIRED {abs(r['days_to_expiry'])}d"
        _add_table_row(full_tbl,
                       [r["hub"], r["site_code"], r["permit_type"],
                        r["expiry_date"], days_str, r["rag"]],
                       bgs=["F5F5F5", "F5F5F5", "FFFFFF", "FFFFFF", rag_bg, rag_bg],
                       bolds=[True, False, False, False, False, True],
                       aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT,
                                WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                                WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER])

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 3: HSE
    # ─────────────────────────────────────────────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, "3. HSE Compliance & Incident Trends")
    inc = hse_report["this_week_counts"]
    avg = hse_report["rolling_avg"]
    _add_body(doc, (
        f"This week: {hse_report['total_this_week']} incident(s)  ·  "
        f"MTC={inc.get('MTC',0)}  FAC={inc.get('FAC',0)}  NM={inc.get('NM',0)}  ·  "
        f"Based on {hse_report['weeks_analyzed']} week(s) of data"
    ))

    _add_heading(doc, "Incident Count vs Rolling Average", level=2)
    trend_tbl = doc.add_table(rows=1, cols=4)
    _table_header_row(trend_tbl, ["Type", "This Week", f"Rolling Avg ({hse_report['weeks_analyzed']}wk)", "vs Avg"],
                      [3.0, 3.0, 4.5, 2.5])
    for t in ("MTC", "FAC", "NM"):
        this = inc.get(t, 0)
        avg_val = avg.get(t, 0.0)
        diff = this - avg_val
        bg = "FFCCCC" if diff > 0.5 else "CCFFCC" if diff < -0.5 else "FFFFFF"
        diff_str = f"+{diff:.1f}" if diff > 0 else f"{diff:.1f}"
        _add_table_row(trend_tbl, [t, str(this), f"{avg_val:.1f}", diff_str],
                       bgs=["F5F5F5", "FFFFFF", "FFFFFF", bg],
                       bolds=[True, False, False, True],
                       aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                                WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER])

    if hse_report["overdue_actions"]:
        _add_heading(doc, f"Overdue Corrective Actions ({hse_report['overdue_count']})", level=2)
        ov_tbl = doc.add_table(rows=1, cols=5)
        _table_header_row(ov_tbl, ["Hub", "Type", "Description", "Due Date", "Days Overdue"],
                          [2.5, 1.5, 7.0, 2.5, 2.5])
        for action in hse_report["overdue_actions"]:
            _add_table_row(ov_tbl,
                           [action["hub"], action["incident_type"],
                            action["description"], action["due_date"],
                            str(action["days_overdue"])],
                           bgs=["F5F5F5", "FFCCCC", "FFFFFF", "FFFFFF", "FFCCCC"],
                           bolds=[True, True, False, False, True],
                           aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                                   WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                                   WD_ALIGN_PARAGRAPH.CENTER])
    else:
        _add_body(doc, "✓ No overdue corrective actions this week.", color=RGBColor(0, 120, 0), bold=True)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 4: Recommended Actions
    # ─────────────────────────────────────────────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, "4. Recommended Actions")
    _add_body(doc, "Auto-generated priority action list. Assign owners before Wednesday standup.")

    actions = []
    for miss in kpi_report["top_misses"][:3]:
        actions.append(("HIGH", f"KPI", miss["hub"],
                         f"{miss['metric']} dropped {abs(miss['delta_pct']):.1f}% WoW — investigate root cause"))
    for item in compliance_report["action_list"][:5]:
        priority = "CRITICAL" if item["rag"] == "RED" else "MEDIUM"
        expired_txt = f"EXPIRED {abs(item['days_to_expiry'])}d ago" if item["days_to_expiry"] < 0 else f"expires in {item['days_to_expiry']}d"
        actions.append((priority, "Compliance", item["hub"],
                         f"{item['permit_type']} {expired_txt} — status: {item['renewal_status']}"))
    for act in hse_report["overdue_actions"][:3]:
        actions.append(("HIGH", "HSE", act["hub"],
                         f"[{act['incident_type']}] {act['description'][:60]} — {act['days_overdue']}d overdue"))

    priority_order = {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2}
    actions.sort(key=lambda x: priority_order.get(x[0], 3))

    act_tbl = doc.add_table(rows=1, cols=4)
    _table_header_row(act_tbl, ["Priority", "Function", "Hub", "Action Required"],
                      [2.0, 2.5, 2.5, 9.0])
    for priority, fn, hub, action in actions:
        bg = {"CRITICAL": "FFCCCC", "HIGH": "FFF0B0", "MEDIUM": "FFFFFF"}.get(priority, "FFFFFF")
        _add_table_row(act_tbl, [priority, fn, hub, action],
                       bgs=[bg, "F5F5F5", "F5F5F5", "FFFFFF"],
                       bolds=[True, False, True, False])

    # ── Footer note ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_r = footer_p.add_run(
        f"Auto-generated by SPX Cowork Automation  ·  {as_of.strftime('%d %b %Y')}  ·  "
        "Not for external distribution"
    )
    footer_r.font.name = "Garamond"
    footer_r.font.size = Pt(8)
    footer_r.font.color.rgb = GREY
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)
